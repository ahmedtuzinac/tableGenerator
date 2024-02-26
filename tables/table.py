import docx

from tables import *
from utils import font

map_of_styles = {
    'font-name': font.set_font_name,
    'font-size': font.set_font_size
}


async def apply_styles(styles: dict, element):
    not_applied_styles = set()

    for style in styles:
        if style not in available_styles:
            not_applied_styles.add(style)
            continue

        kwargs = {
            f'{style}': styles[style],
            'element': element
        }
        await map_styles(style, kwargs)


async def map_styles(style: str, kwargs):
    func = map_of_styles.get(style)
    await func(**kwargs)


async def create_table():
    doc = docx.Document()
    await apply_styles(document['styles'], doc)

    for table in tables['content']:
        try:
            number_of_columns = len(table['rows'][0])
            # TODO: make def which returns max len(row) :)
        except IndexError:
            continue

        added_table = doc.add_table(rows=1, cols=number_of_columns)
        for row in table['rows']:
            new_row = added_table.add_row()
            for idx, cell in enumerate(new_row.cells):
                try:
                    cell.text = row[idx]
                except IndexError:
                    cell.text = "null"

        added_table.rows[0].cells[0].merge(added_table.rows[0].cells[-1])
        header_text = table.get('header').get('content')
        added_table.rows[0].cells[0].text = header_text if header_text else 'Header'

        if tables.get('style'):
            await apply_styles(tables['style'], added_table)

        if table.get('style'):
            await apply_styles(table['style'], added_table)

        if table.get('header').get('style'):
            await apply_styles(table['header']['style'], added_table.rows[0])

        doc.add_paragraph()

    doc_name = document.get('name') if document.get('name') else 'output'
    doc.save(f'{file_path}/outputs/{doc_name}.docx')
