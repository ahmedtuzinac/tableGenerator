import docx
from docx.shared import Pt


async def set_font_name(**kwargs):
    element = kwargs.get('element')
    font_name = kwargs.get('font-name')

    if isinstance(element, docx.document.Document):
        normal_style = element.styles['Normal']
        normal_style.font.name = font_name
        return

    if isinstance(element, docx.table.Table):
        for row in element.rows:
            for cell in row.cells:
                cell.text = 'test'
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name

    if isinstance(element, docx.table._Row):
        for cell in element.cells:
            cell.text = 'test'
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
        return


async def set_font_size(**kwargs):
    element = kwargs.get('element')
    font_size = Pt(kwargs.get('font-size'))

    if isinstance(element, docx.document.Document):
        normal_style = element.styles['Normal']
        normal_style.font.size = font_size
        return

    if isinstance(element, docx.table.Table):
        for row in element.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = font_size
        return

    if isinstance(element, docx.table._Row):
        for cell in element.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size
        return
